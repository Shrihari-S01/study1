"""Word document generation."""

import asyncio
from pathlib import Path

from docx import Document
from docx.shared import Inches

from app.core.config import get_settings
from app.models.auction import Auction


class WordGenerator:
    """Generate a Word file for an auction notice."""

    def __init__(self) -> None:
        self.settings = get_settings()

    async def generate(self, auction: Auction) -> Path:
        """Generate a .docx document for the auction."""
        self.settings.word_output_dir.mkdir(parents=True, exist_ok=True)
        path = self.settings.word_output_dir / f"{auction.listing_id}.docx"
        await asyncio.to_thread(self._generate_sync, auction, path)
        return path

    def _generate_sync(self, auction: Auction, path: Path) -> None:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        document.add_heading("Auction Notice Details", level=1)
        document.add_paragraph(f"Listing ID: {auction.listing_id}")

        rows = [
            ("Bank Name", auction.bank_name),
            ("Borrower Name", auction.borrower_name),
            ("Loan Number", auction.loan_number),
            ("Auction Type", auction.auction_type),
            ("Property Type", auction.property_type),
            ("Property Category", auction.property_category),
            ("Asset Category", auction.asset_category),
            ("Movable/Immovable", auction.movable_immovable),
            ("Possession Type", auction.possession_type),
            ("Reserve Price", auction.reserve_price),
            ("EMD", auction.emd),
            ("Demand Notice Date", auction.demand_notice_date),
            ("Symbolic Possession Date", auction.symbolic_possession_date),
            ("Auction Date", auction.auction_date),
            ("Property Address", auction.property_address),
            ("District", auction.district),
            ("State", auction.state),
            ("Beneficiary Bank", auction.beneficiary_bank),
            ("IFSC", auction.ifsc),
            ("Contact Person", auction.contact_person),
            ("Contact Number", auction.contact_number),
            ("Website", auction.website),
        ]

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value or ""

        document.add_heading("Description", level=2)
        document.add_paragraph(auction.description or "")

        document.add_heading("4W Analysis", level=2)
        for label, value in (
            ("WHO", auction.who),
            ("WHOM", auction.whom),
            ("WHERE", auction.where_location),
            ("WHEN", auction.when_details),
        ):
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value or "")

        document.add_heading("Summary", level=2)
        document.add_paragraph(auction.summary or "")
        document.add_paragraph(f"Confidence Score: {auction.confidence_score:.2f}")
        document.save(path)

